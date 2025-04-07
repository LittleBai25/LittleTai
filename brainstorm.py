import streamlit as st
import os
import tempfile
import re
from pathlib import Path
import io

# 导入基本依赖
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from PIL import Image
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False

# 导入 LangChain 相关库
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.callbacks.streamlit import StreamlitCallbackHandler

# 页面配置
st.set_page_config(
    page_title="脑暴助理",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 设置API客户端
def get_langchain_llm(model_type="simplify", stream=False, st_container=None):
    """根据不同的模型类型设置API客户端"""
    # 使用OpenRouter API
    api_base = "https://openrouter.ai/api/v1"
    
    if model_type == "simplify":
        # 素材分析使用的API密钥和模型
        api_key = st.secrets.get("OPENROUTER_API_KEY_SIMPLIFY", "")
        model_name = "deepseek/deepseek-chat-v3-0324:free"  # 使用deepseek模型
        temperature = 0.1  # 降低温度以获得更稳定的输出
    else:  # analysis
        # 脑暴报告使用的API密钥和模型
        api_key = st.secrets.get("OPENROUTER_API_KEY_ANALYSIS", "")
        model_name = "deepseek/deepseek-chat-v3-0324:free"  # 使用deepseek模型
        temperature = 0.3  # 降低温度以获得更稳定的输出
        
    # 检查API密钥是否为空
    if not api_key:
        st.error(f"{'素材分析' if model_type == 'simplify' else '脑暴报告'} API密钥未设置！请在secrets.toml中配置。")
        st.stop()
    
    # 设置回调处理器
    callbacks = None
    if stream and st_container:
        callbacks = [StreamlitCallbackHandler(st_container)]
    
    try:
        # 创建LangChain LLM客户端 - 完全移除所有限制
        llm = OpenAI(
            model_name=model_name,
            openai_api_key=api_key,
            openai_api_base=api_base,
            streaming=stream,
            temperature=temperature,
            callbacks=callbacks,
            request_timeout=120,  # 增加超时时间到120秒
            max_retries=5,  # 增加重试次数
            presence_penalty=0,  # 移除存在惩罚
            frequency_penalty=0,  # 移除频率惩罚
            max_tokens=None,  # 明确设置为None，表示不限制token
            top_p=1.0,  # 允许所有可能的token
            best_of=1,  # 使用最佳结果
            n=1,  # 只生成一个结果
            logit_bias={},  # 不设置任何token偏见
            stop=None  # 不设置停止条件
        )
        
        # 测试API连接 - 使用generate方法而不是predict
        test_prompt = "Hello, this is a test."
        try:
            test_response = llm.generate([test_prompt])
            if not test_response or not test_response.generations:
                raise Exception("API测试响应为空")
            st.write("API连接测试成功")
        except Exception as e:
            st.error(f"API连接测试失败: {str(e)}")
            st.write("正在尝试使用备用配置...")
            # 尝试使用备用配置
            llm = OpenAI(
                model_name=model_name,
                openai_api_key=api_key,
                openai_api_base=api_base,
                streaming=False,  # 禁用流式输出
                temperature=temperature,
                request_timeout=180,  # 进一步增加超时时间
                max_retries=10,  # 增加重试次数
                max_tokens=None
            )
        
        return llm
    except Exception as e:
        st.error(f"创建LLM客户端时出错: {str(e)}")
        st.write("错误详情：")
        st.write(str(e))
        st.stop()

# 文件处理函数
def process_file(file_path, file_type):
    """处理不同类型的文件并返回内容"""
    try:
        # 检查文件是否存在并有内容
        if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
            return f"警告: 文件 {os.path.basename(file_path)} 为空或不存在"
            
        if file_type == "docx" and DOCX_SUPPORT:
            try:
                # 使用python-docx处理
                doc = docx.Document(file_path)
                content_parts = []
                
                # 提取段落文本
                for para in doc.paragraphs:
                    if para.text.strip():
                        # 增强格式处理，保留更多格式信息
                        para_text = ""
                        for run in para.runs:
                            text = run.text.strip()
                            if not text:
                                continue
                                
                            # 处理加粗
                            if run.bold:
                                text = f"**{text}**"
                            # 处理斜体
                            if run.italic:
                                text = f"*{text}*"
                            # 处理下划线
                            if run.underline:
                                text = f"__{text}__"
                            # 处理字体大小
                            if run.font.size:
                                size = run.font.size.pt if run.font.size.pt else 11
                                if size > 11:
                                    text = f"# {text}" if size > 14 else f"## {text}"
                            
                            para_text += text + " "
                        
                        # 清理多余空格并添加段落
                        if para_text.strip():
                            content_parts.append(para_text.strip())
                
                # 提取表格内容
                for table_idx, table in enumerate(doc.tables):
                    if len(table.rows) == 0:
                        continue
                    
                    # 添加表格标记
                    content_parts.append(f"\n## 表格 {table_idx+1}")
                    
                    # 判断表格类型和结构
                    is_questionnaire = False
                    if len(table.rows) > 1 and len(table.rows[0].cells) > 0:
                        # 检查第一行是否可能是表头
                        header_row = [cell.text.strip() for cell in table.rows[0].cells]
                        is_questionnaire = any("问题" in cell or "题" in cell for cell in header_row) or len(header_row) >= 2
                    
                    if is_questionnaire:
                        # 特殊处理问卷表格，按行分组
                        headers = []
                        for cell in table.rows[0].cells:
                            headers.append(cell.text.strip())
                        
                        # 处理内容行
                        for row_idx in range(1, len(table.rows)):
                            row_content = []
                            for col_idx, cell in enumerate(table.rows[row_idx].cells):
                                cell_text = cell.text.strip()
                                # 如果有表头且内容不为空，关联显示
                                if cell_text and col_idx < len(headers) and headers[col_idx]:
                                    row_content.append(f"{headers[col_idx]}: {cell_text}")
                                elif cell_text:
                                    row_content.append(cell_text)
                            
                            # 只添加非空内容
                            if row_content:
                                content_parts.append(" | ".join(row_content))
                    else:
                        # 常规表格处理
                        for row in table.rows:
                            row_texts = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    # 替换可能导致格式问题的字符
                                    cell_text = cell_text.replace('\n', ' ').replace('|', '/')
                                    row_texts.append(cell_text)
                            
                            # 只添加非空行
                            if row_texts:
                                content_parts.append(" | ".join(row_texts))
                
                # 合并所有内容，添加适当的换行
                content = "\n\n".join(content_parts)
                
                # 记录日志
                st.write(f"从DOCX文件 {os.path.basename(file_path)} 读取了 {len(content)} 字符")
                    
                # 后处理，清理可能的重复内容和格式标记
                content = content.replace('{.mark}', '').replace('{.underline}', '')
                
                # 确保内容不为空
                if not content.strip():
                    return f"警告: 文件 {os.path.basename(file_path)} 内容为空"
                
                return content
            except Exception as e:
                error_msg = f"读取DOCX文件时出错: {str(e)}"
                st.error(error_msg)
                return error_msg
                
        elif file_type == "pdf" and PDF_SUPPORT:
            try:
                pdf_reader = PdfReader(file_path)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                st.write(f"从PDF文件 {os.path.basename(file_path)} 读取了 {len(text)} 字符")
                return text
            except Exception as e:
                error_msg = f"读取PDF文件时出错: {str(e)}"
                st.error(error_msg)
                return error_msg
        elif file_type in ["jpg", "jpeg", "png"] and IMAGE_SUPPORT:
            # 简单记录图像信息，而不进行OCR
            try:
                image = Image.open(file_path)
                width, height = image.size
                info = f"[图像文件，尺寸: {width}x{height}，类型: {image.format}。请在分析时考虑此图像可能包含的视觉内容。]"
                st.write(f"处理图像文件: {os.path.basename(file_path)}")
                return info
            except Exception as e:
                error_msg = f"处理图像文件时出错: {str(e)}"
                st.error(error_msg)
                return error_msg
        else:
            # 尝试作为文本文件读取
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    st.write(f"从文本文件 {os.path.basename(file_path)} 读取了 {len(content)} 字符")
                    return content
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='ignore')
                        st.write(f"从二进制文件 {os.path.basename(file_path)} 读取了 {len(content)} 字符")
                        return content
                except Exception as e:
                    error_msg = f"以二进制模式读取文件时出错: {str(e)}"
                    st.error(error_msg)
                    return error_msg
            except Exception as e:
                error_msg = f"读取文本文件时出错: {str(e)}"
                st.error(error_msg)
                return error_msg
    except Exception as e:
        error_msg = f"处理文件时出错: {str(e)}"
        st.error(error_msg)
        return error_msg

# 简化文件内容
def chunk_content(content, chunk_size=8000):
    """将内容分块处理"""
    words = content.split()
    chunks = []
    current_chunk = []
    current_size = 0
    
    for word in words:
        word_size = len(word) + 1  # +1 for space
        if current_size + word_size > chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_size = word_size
        else:
            current_chunk.append(word)
            current_size += word_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def simplify_content(content, direction, st_container=None):
    """使用AI简化上传的文件内容"""
    try:
        # 记录日志，确认内容长度
        st.write(f"准备分析的内容总长度: {len(content)} 字符")
        st.write("内容预览（前500字符）:")
        st.write(content[:500])
        
        # 检查内容是否有效
        if not content or len(content.strip()) < 10:
            st.error("文档内容过短或为空")
            return "文档内容过短或为空，请检查上传的文件是否正确"
            
        # 获取API客户端 - 使用带有备用方案的流式输出
        llm = get_langchain_llm("simplify", stream=True, st_container=st_container)
        
        # 从会话状态获取提示词
        backstory = st.session_state.material_backstory_prompt
        task = st.session_state.material_task_prompt
        output_format = st.session_state.material_output_prompt
        
        # 清理文本，移除可能导致问题的特殊字符
        clean_content = content.replace('{.mark}', '').replace('{.underline}', '')
        clean_content = clean_content.replace('\x00', '')  # 移除空字符
        clean_content = re.sub(r'\s+', ' ', clean_content)  # 规范化空白字符
        
        # 记录清理后的内容长度
        st.write(f"清理后的内容长度: {len(clean_content)} 字符")
        st.write("清理后的内容预览（前500字符）:")
        st.write(clean_content[:500])
        
        # 构建完整的提示词模板
        template = f"""{backstory}

{task}

{output_format}

重要要求:
1. 你必须详细分析文档的每一部分内容，不要遗漏任何细节
2. 提取所有与研究方向相关的信息，包括隐含的信息
3. 保持原文的层次结构，使用清晰的标题和列表
4. 如果文档包含表格，必须完整保留表格的结构和内容
5. 如果文档包含图片，必须详细描述图片的内容和位置
6. 输出必须包含以下部分：
   - 文档概述（详细说明文档的主要内容和结构）
   - 关键发现（列出所有重要的发现和见解）
   - 具体细节（详细分析每个重要部分）
   - 相关表格内容（完整保留和解释所有表格）
   - 相关图片描述（详细描述所有图片及其重要性）
   - 总结和建议（提供全面的总结和具体建议）
7. 每个部分都必须详细展开，提供充分的解释和分析
8. 不要限制输出长度，确保完整分析所有内容
9. 不要遗漏任何重要信息，包括看似次要的细节
10. 如果遇到不确定的内容，请明确标注并说明原因

研究方向: {direction}

文档内容:
{clean_content}

请按照上述要求生成详细的分析结果。输出必须包含所有要求的部分，并且每个部分都要详细展开。不要限制输出长度，确保完整分析所有内容。"""
        
        prompt = PromptTemplate(
            template=template,
            input_variables=["direction", "clean_content"]
        )
        
        # 创建LLMChain
        chain = LLMChain(llm=llm, prompt=prompt)
        
        # 执行链 - 使用更明确的参数名称
        with st.spinner("正在分析文档内容..."):
            try:
                # 记录API调用开始
                st.write("开始API调用...")
                result = chain.run(direction=direction, clean_content=clean_content)
                st.write("API调用成功，正在处理结果...")
            except Exception as e:
                st.error(f"API调用失败: {str(e)}")
                st.write("错误详情：")
                st.write(str(e))
                st.write("正在尝试使用备用模型...")
                # 尝试使用备用模型
                llm = get_langchain_llm("simplify", stream=False, st_container=st_container)
                chain = LLMChain(llm=llm, prompt=prompt)
                try:
                    result = chain.run(direction=direction, clean_content=clean_content)
                    st.write("备用模型调用成功")
                except Exception as e2:
                    st.error(f"备用模型调用也失败: {str(e2)}")
                    st.write("错误详情：")
                    st.write(str(e2))
                    return "AI分析失败。请检查API密钥是否正确，或稍后重试。"
        
        # 检查结果是否有效
        if not result or len(result.strip()) < 10:
            st.error("AI分析未能生成有效结果")
            st.write("可能的原因：")
            st.write("1. 文档内容可能过于复杂或格式特殊")
            st.write("2. 研究方向描述可能不够明确")
            st.write("3. API调用可能出现了问题")
            st.write("4. 文档内容可能包含特殊字符或格式")
            st.write("5. API密钥可能无效或过期")
            st.write("6. 网络连接可能不稳定")
            return "AI分析未能生成有效结果。请检查文档内容是否相关，或调整提示词设置。"
        
        # 记录生成结果的长度
        st.write(f"生成的分析结果长度: {len(result)} 字符")
        st.write("结果预览（前500字符）:")
        st.write(result[:500])
        
        return result
    except Exception as e:
        st.error(f"分析过程中发生错误: {str(e)}")
        st.write("错误详情：")
        st.write(str(e))
        return f"分析过程中发生错误: {str(e)}"

# 生成分析报告
def generate_analysis(simplified_content, direction, st_container=None):
    """使用AI生成分析报告"""
    # 使用流式输出
    llm = get_langchain_llm("analysis", stream=True, st_container=st_container)
    
    try:
        # 检查简化内容是否有效
        if not simplified_content or len(simplified_content.strip()) < 100:
            return "无法生成报告，因为文档分析阶段未能产生足够深入的内容。请返回上一步重试，调整研究方向或上传更相关的文档。"
            
        # 从会话状态获取提示词
        backstory = st.session_state.brainstorm_backstory_prompt
        task = st.session_state.brainstorm_task_prompt
        output_format = st.session_state.brainstorm_output_prompt
        
        # 增强提示模板的明确性和结构
        template = f"""{backstory}

{task}

{output_format}

重要要求:
1. 基于提供的分析结果，生成一份详尽、实用的报告
2. 报告必须与研究方向"{direction}"紧密结合
3. 提供具体的、可实施的策略和方案
4. 包含清晰的结构和小标题
5. 内容必须具备原创性和创新性

研究方向: {direction}

分析结果:
{simplified_content}

请生成一份全面的申请策略和提升方案报告，确保包含明确的小标题和结构化内容。"""
        
        prompt = PromptTemplate(
            template=template,
            input_variables=["direction", "simplified_content"]
        )
        
        # 创建LLMChain
        chain = LLMChain(llm=llm, prompt=prompt)
        
        # 执行链
        result = chain.run(direction=direction, simplified_content=simplified_content)
        
        # 如果返回为空或过短，提供更明确的错误信息
        if not result or len(result.strip()) < 200:
            return "生成报告失败。AI未能生成有意义的内容，可能是因为分析内容不够详细或研究方向过于模糊。请调整提示词设置或返回上一步提供更充分的信息。"
        
        return result
    except Exception as e:
        st.error(f"生成报告时出错: {str(e)}")
        return f"生成报告时出错: {str(e)}"

# 保存提示词函数
def save_prompts():
    """保存当前的提示词到会话状态"""
    # 保存素材分析提示词
    st.session_state['material_backstory_prompt'] = st.session_state.material_backstory_prompt_input
    st.session_state['material_task_prompt'] = st.session_state.material_task_prompt_input
    st.session_state['material_output_prompt'] = st.session_state.material_output_prompt_input
    
    # 保存脑暴报告提示词
    st.session_state['brainstorm_backstory_prompt'] = st.session_state.brainstorm_backstory_prompt_input
    st.session_state['brainstorm_task_prompt'] = st.session_state.brainstorm_task_prompt_input
    st.session_state['brainstorm_output_prompt'] = st.session_state.brainstorm_output_prompt_input
    
    st.success("提示词已保存!")

# 初始化会话状态变量
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = []
if 'direction' not in st.session_state:
    st.session_state.direction = ""
if 'simplified_content' not in st.session_state:
    st.session_state.simplified_content = ""
if 'analysis_report' not in st.session_state:
    st.session_state.analysis_report = ""
if 'show_analysis_section' not in st.session_state:
    st.session_state.show_analysis_section = False

# 素材分析提示词初始化
if 'material_backstory_prompt' not in st.session_state:
    st.session_state.material_backstory_prompt = "你是一个专业的素材内容分析助手。"
if 'material_task_prompt' not in st.session_state:
    st.session_state.material_task_prompt = "请根据用户的方向，提取并分析文档中的关键信息。"
if 'material_output_prompt' not in st.session_state:
    st.session_state.material_output_prompt = "以清晰的要点形式组织输出内容，突出关键信息和见解。"

# 脑暴报告提示词初始化
if 'brainstorm_backstory_prompt' not in st.session_state:
    st.session_state.brainstorm_backstory_prompt = "你是一个专业的头脑风暴报告生成助手。"
if 'brainstorm_task_prompt' not in st.session_state:
    st.session_state.brainstorm_task_prompt = "你的任务是根据素材分析内容和用户的研究方向，生成一份创新的头脑风暴报告。"
if 'brainstorm_output_prompt' not in st.session_state:
    st.session_state.brainstorm_output_prompt = "报告应包括关键发现、创新思路、潜在机会和具体建议，格式清晰易读。"

# 创建两个标签页
tab1, tab2 = st.tabs(["脑暴助理", "管理员设置"])

# 用户界面标签页
with tab1:
    st.title("🧠 脑暴助理")
    st.markdown("欢迎使用脑暴助理！上传您的文件，输入研究方向，获取专业分析报告。")

    # 第一步：上传文件和输入方向
    st.header("第一步：上传文件和输入研究方向")
    
    uploaded_files = st.file_uploader("上传文件（支持DOC, DOCX, PDF, JPG, PNG, TXT）", 
                                     type=['doc', 'docx', 'pdf', 'jpg', 'jpeg', 'png', 'txt'], 
                                     accept_multiple_files=True)
    
    direction = st.text_area("请输入您的研究方向", 
                             height=100, 
                             help="详细描述您的研究方向，帮助AI更好地理解您的需求")
    
    if st.button("开始素材分析", disabled=not uploaded_files or not direction):
        # 保存上传的文件到临时目录
        temp_dir = tempfile.mkdtemp()
        file_paths = []
        
        # 保存文件并添加到处理列表
        for file in uploaded_files:
            # 使用安全的文件名，移除特殊字符
            safe_filename = re.sub(r'[^\w\-\.]', '_', file.name)
            file_path = os.path.join(temp_dir, safe_filename)
            with open(file_path, "wb") as f:
                f.write(file.getbuffer())
            file_paths.append(file_path)
            st.write(f"保存文件: {file.name} -> {file_path}, 大小: {len(file.getbuffer())} 字节")
        
        # 确保立即保存方向信息到会话状态
        st.session_state.uploaded_files = file_paths
        st.session_state.direction = direction
        
        # 处理上传的文件内容，逐个处理每个文件并收集内容
        all_content = ""
        for file_path in file_paths:
            file_ext = Path(file_path).suffix.lower().replace(".", "")
            
            # 使用process_file函数提取文件内容，并添加到all_content
            content = process_file(file_path, file_ext)
            file_name = os.path.basename(file_path)
            
            all_content += f"\n\n===== 文件: {file_name} =====\n\n{content}"
            
        # 验证文件内容不为空
        if not all_content or len(all_content.strip()) < 50:
            st.error("❌ 文件内容似乎为空或过短。请确保上传了有效的文件。")
            st.stop()
        
        # 创建一个容器用于流式输出
        analysis_container = st.empty()
        
        # 简化内容
        with st.spinner("正在分析素材..."):
            # 调用AI分析内容
            simplified = simplify_content(all_content, direction, st_container=analysis_container)
            
            # 确保立即保存简化内容到会话状态
            st.session_state.simplified_content = simplified
            st.session_state.show_analysis_section = True
        
        # 显示结果
        st.subheader("素材分析结果")
        st.markdown(simplified)
    
    # 第二步：生成头脑风暴辅助报告
    if st.session_state.show_analysis_section or st.session_state.simplified_content:
        st.header("第二步：生成头脑风暴辅助报告")
        
        # 每次UI渲染时都确保研究方向同步更新
        if direction and direction != st.session_state.direction:
            st.session_state.direction = direction

        if st.button("生成脑暴报告", disabled=not (st.session_state.simplified_content and st.session_state.direction)):
            # 使用已经生成的简化内容和研究方向
            
            # 创建一个容器用于流式输出
            report_container = st.empty()
            
            # 生成分析报告
            with st.spinner("正在生成脑暴报告..."):
                report = generate_analysis(st.session_state.simplified_content, st.session_state.direction, st_container=report_container)
                st.session_state.analysis_report = report
            
            # 显示结果
            st.subheader("脑暴报告")
            st.markdown(report)

# 管理员设置标签页
with tab2:
    st.title("🔧 管理员设置")
    st.markdown("配置AI提示词")
    
    # 素材分析提示词设置
    st.header("素材分析提示词设置")
    
    st.subheader("素材分析 - Backstory")
    material_backstory_prompt = st.text_area("素材分析AI背景设定", 
                                   value=st.session_state.material_backstory_prompt,
                                   height=100,
                                   key="material_backstory_prompt_input",
                                   help="设定素材分析AI的角色和背景")
    
    st.subheader("素材分析 - Task Description")
    material_task_prompt = st.text_area("素材分析任务描述", 
                              value=st.session_state.material_task_prompt,
                              height=100,
                              key="material_task_prompt_input",
                              help="描述素材分析AI需要执行的具体任务")
    
    st.subheader("素材分析 - Output Format")
    material_output_prompt = st.text_area("素材分析输出格式", 
                                value=st.session_state.material_output_prompt,
                                height=100,
                                key="material_output_prompt_input",
                                help="指定素材分析AI输出的格式和风格")
    
    # 脑暴报告提示词设置
    st.header("脑暴报告提示词设置")
    
    st.subheader("脑暴报告 - Backstory")
    brainstorm_backstory_prompt = st.text_area("脑暴报告AI背景设定", 
                                   value=st.session_state.brainstorm_backstory_prompt,
                                   height=100,
                                   key="brainstorm_backstory_prompt_input",
                                   help="设定脑暴报告AI的角色和背景")
    
    st.subheader("脑暴报告 - Task Description")
    brainstorm_task_prompt = st.text_area("脑暴报告任务描述", 
                              value=st.session_state.brainstorm_task_prompt,
                              height=100,
                              key="brainstorm_task_prompt_input",
                              help="描述脑暴报告AI需要执行的具体任务")
    
    st.subheader("脑暴报告 - Output Format")
    brainstorm_output_prompt = st.text_area("脑暴报告输出格式", 
                                value=st.session_state.brainstorm_output_prompt,
                                height=100,
                                key="brainstorm_output_prompt_input",
                                help="指定脑暴报告AI输出的格式和风格")
    
    if st.button("保存提示词设置"):
        save_prompts()

# 添加页脚
st.markdown("---")
st.markdown("© 2025 脑暴助理 | 由Streamlit、LangChain和OpenRouter提供支持")
